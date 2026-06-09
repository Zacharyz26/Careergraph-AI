from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

import fitz
from docx import Document


@dataclass(frozen=True)
class ParsedDocument:
    extracted_text: str
    file_type: str
    page_count: int | None = None


class DocumentParserError(ValueError):
    """Base error for documents that cannot be accepted or parsed."""


class UnsupportedDocumentTypeError(DocumentParserError):
    pass


class EmptyDocumentError(DocumentParserError):
    pass


class InvalidDocumentError(DocumentParserError):
    pass


class NoExtractableTextError(DocumentParserError):
    pass


class DocumentParser:
    supported_extensions = {".pdf", ".docx"}

    def parse(self, *, filename: str, content: bytes) -> ParsedDocument:
        extension = Path(filename).suffix.lower()

        if extension not in self.supported_extensions:
            raise UnsupportedDocumentTypeError(
                "Unsupported file type. Only PDF and DOCX files are accepted."
            )
        if not content:
            raise EmptyDocumentError("The uploaded file is empty.")

        if extension == ".pdf":
            parsed = self._parse_pdf(content)
        else:
            parsed = self._parse_docx(content)

        if not parsed.extracted_text.strip():
            raise NoExtractableTextError(
                "The uploaded file contains no extractable text."
            )
        return parsed

    def _parse_pdf(self, content: bytes) -> ParsedDocument:
        try:
            with fitz.open(stream=content, filetype="pdf") as document:
                page_count = document.page_count
                page_texts = []
                for page in document:
                    page_text = page.get_text("text").strip()
                    if page_text:
                        page_texts.append(page_text)
                text = "\n\n".join(page_texts)
        except Exception as exc:
            raise InvalidDocumentError(
                "The uploaded PDF is invalid or corrupted."
            ) from exc

        return ParsedDocument(
            extracted_text=text.strip(),
            file_type="pdf",
            page_count=page_count,
        )

    def _parse_docx(self, content: bytes) -> ParsedDocument:
        try:
            document = Document(BytesIO(content))
        except Exception as exc:
            raise InvalidDocumentError(
                "The uploaded DOCX file is invalid or corrupted."
            ) from exc

        text_parts = [
            paragraph.text.strip()
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]
        text_parts.extend(
            cell.text.strip()
            for table in document.tables
            for row in table.rows
            for cell in row.cells
            if cell.text.strip()
        )

        return ParsedDocument(
            extracted_text="\n".join(text_parts).strip(),
            file_type="docx",
        )
