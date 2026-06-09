from io import BytesIO

import fitz
import pytest
from docx import Document

from app.services.document_parser import (
    DocumentParser,
    EmptyDocumentError,
    NoExtractableTextError,
    UnsupportedDocumentTypeError,
)


def test_extracts_text_and_page_count_from_pdf() -> None:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "CareerGraph PDF resume")
    content = document.tobytes()
    document.close()

    result = DocumentParser().parse(filename="resume.pdf", content=content)

    assert result.file_type == "pdf"
    assert result.page_count == 1
    assert "CareerGraph PDF resume" in result.extracted_text


def test_extracts_paragraph_and_table_text_from_docx() -> None:
    document = Document()
    document.add_paragraph("CareerGraph DOCX resume")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Python"
    buffer = BytesIO()
    document.save(buffer)

    result = DocumentParser().parse(
        filename="resume.DOCX",
        content=buffer.getvalue(),
    )

    assert result.file_type == "docx"
    assert result.page_count is None
    assert "CareerGraph DOCX resume" in result.extracted_text
    assert "Python" in result.extracted_text


def test_rejects_unsupported_file_type() -> None:
    with pytest.raises(UnsupportedDocumentTypeError, match="Only PDF and DOCX"):
        DocumentParser().parse(filename="resume.txt", content=b"resume")


def test_rejects_empty_file() -> None:
    with pytest.raises(EmptyDocumentError, match="empty"):
        DocumentParser().parse(filename="resume.pdf", content=b"")


def test_rejects_document_without_extractable_text() -> None:
    document = Document()
    buffer = BytesIO()
    document.save(buffer)

    with pytest.raises(NoExtractableTextError, match="no extractable text"):
        DocumentParser().parse(
            filename="resume.docx",
            content=buffer.getvalue(),
        )
